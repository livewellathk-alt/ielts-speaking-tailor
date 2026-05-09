from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def render_markdown(payload: dict[str, Any]) -> str:
    answers = payload["answers"]
    lines = [
        "# IELTS Speaking Tailor",
        "",
        "## Style Guide",
        "",
        f"- Student voice: {payload.get('style_guide', {}).get('student_voice', 'not specified')}",
        "",
    ]
    word_targets = payload.get("word_targets", {})
    if word_targets:
        part1 = word_targets.get("part1", {})
        part2 = word_targets.get("part2", {})
        part3 = word_targets.get("part3", {})
        lines.extend(
            [
                "## Timing Targets",
                "",
                f"- Part 1: about {part1.get('seconds')} seconds / {part1.get('words')} words",
                (
                    f"- Part 2: {part2.get('min_seconds')}-{part2.get('max_seconds')} seconds / "
                    f"{part2.get('min_words')}-{part2.get('max_words')} words"
                ),
                f"- Part 3: about {part3.get('seconds')} seconds / {part3.get('words')} words",
                "",
            ]
        )
    lines.extend(["## Umbrella Story Index", ""])
    story_map: dict[str, list[str]] = {}
    for block in answers.get("part2_blocks", []):
        story_map.setdefault(_umbrella_story_id(block.get("umbrella_story", "story_general")), []).append(block.get("title_zh", block.get("block_id", "")))
    for story_id, titles in sorted(story_map.items()):
        lines.append(f"- `{story_id}`: {', '.join(title for title in titles if title)}")
    lines.extend(["", "## Part 1", ""])
    for answer in answers.get("part1", []):
        lines.extend(_answer_section(f"### {answer.get('question', answer.get('question_id', 'Question'))}", answer))
    lines.extend(["", "## Part 2 and Related Part 3", ""])
    for block in answers.get("part2_blocks", []):
        title = block.get("title_zh") or block.get("block_id", "Part 2")
        lines.extend(_answer_section(f"### Part 2: {title}", block, prompt=block.get("part2_prompt")))
        if block.get("part3"):
            lines.extend(["", "#### Part 3", ""])
        for p3 in block.get("part3", []):
            lines.extend(_answer_section(f"##### {p3.get('question', p3.get('question_id', 'Question'))}", p3))
    return "\n".join(lines).strip() + "\n"


def render_outputs(payload: dict[str, Any], *, output_dir: str | Path, basename: str = "ielts_speaking_answers") -> dict[str, Path]:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    markdown_path = output / f"{basename}.md"
    docx_path = output / f"{basename}.docx"
    markdown = render_markdown(payload)
    markdown_path.write_text(markdown, encoding="utf-8")
    _write_docx(markdown, docx_path)
    return {"markdown": markdown_path, "docx": docx_path}


def _answer_section(title: str, answer: dict[str, Any], prompt: str | None = None) -> list[str]:
    lines = [title, ""]
    if prompt:
        lines.extend([f"Prompt: {prompt}", ""])
    if answer.get("framework"):
        lines.extend([f"Framework: {answer['framework']}", ""])
    if answer.get("answer_en"):
        lines.extend([f"English: {answer['answer_en']}", ""])
    if answer.get("answer_zh"):
        lines.extend([f"中文: {answer['answer_zh']}", ""])
    memory_cues = _memory_cues(answer.get("memory_cues"))
    if memory_cues:
        lines.extend([f"Memory cues: {', '.join(memory_cues)}", ""])
    if answer.get("umbrella_story"):
        lines.extend([f"Umbrella story: `{_umbrella_story_id(answer['umbrella_story'])}`", ""])
    return lines


def _umbrella_story_id(value: Any) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        return str(value.get("id") or value.get("title") or value.get("story") or "story_general")
    return str(value or "story_general")


def _memory_cues(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        normalized = value
        for separator in ["，", "、", ";", "；", "\n"]:
            normalized = normalized.replace(separator, ",")
        return [item.strip() for item in normalized.split(",") if item.strip()]
    return []


def _write_docx(markdown: str, path: Path) -> None:
    document = Document()
    for line in markdown.splitlines():
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            document.add_heading(line[5:], level=3)
        elif line.startswith("##### "):
            document.add_heading(line[6:], level=4)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    document.save(path)
